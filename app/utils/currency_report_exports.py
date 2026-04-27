from __future__ import annotations

from io import BytesIO

import pandas as pd
import streamlit as st

from app.utils.currency_utils import CURRENCY_INFO


_ZERO_DECIMAL_CURRENCIES = {
    "JPY", "KRW", "VND", "IDR", "UGX", "TZS", "PYG",
    "BIF", "CLP", "GNF", "MGA", "RWF", "XOF", "XAF",
    "XPF", "IRR", "SYP", "YER", "NGN", "HUF",
}


def _currency_symbol(currency_code: str) -> str:
    info = CURRENCY_INFO.get(currency_code.upper())
    return info[1] if info else f"{currency_code.upper()} "


def _currency_decimals(currency_code: str) -> int:
    return 0 if currency_code.upper() in _ZERO_DECIMAL_CURRENCIES else 2


def _format_currency(amount: float, currency_code: str) -> str:
    decimals = _currency_decimals(currency_code)
    code = currency_code.upper()
    symbol = _currency_symbol(code)
    pdf_safe_prefix = symbol if code in {"USD", "GBP"} else f"{code} "
    if decimals == 0:
        return f"{pdf_safe_prefix}{amount:,.0f}"
    return f"{pdf_safe_prefix}{amount:,.2f}"


def _get_rate(rate_data: dict, target_currency: str) -> float:
    rates = (rate_data or {}).get("rates", {})
    return float(rates.get(target_currency.upper(), 1.0))


def _convert_amount(usd_amount: float | int | None, target_currency: str, rate_data: dict) -> float | None:
    if usd_amount is None:
        return None
    return float(usd_amount) * _get_rate(rate_data, target_currency)


def _summary_header_lines(target_currency: str, rate_data: dict) -> tuple[str, ...]:
    source = (rate_data or {}).get("source", "cached")
    rate = _get_rate(rate_data, target_currency)
    if target_currency.upper() == "USD":
        return (
            "Report currency: USD",
            "This companion export keeps salary values in the app's default report currency.",
        )
    return (
        f"Report currency: {target_currency.upper()}",
        f"Applied exchange rate: 1 USD = {_format_currency(rate, target_currency)} ({source})",
        "This companion export mirrors the current result with salary values shown in the selected currency.",
    )


def _clean_df(df: pd.DataFrame) -> pd.DataFrame:
    cleaned = df.copy().fillna("")
    for col in cleaned.columns:
        if pd.api.types.is_float_dtype(cleaned[col]):
            cleaned[col] = cleaned[col].round(2)
    return cleaned


@st.cache_resource(show_spinner=False)
def _get_pdf_font_names() -> tuple[str, str]:
    """
    Return (regular_font_name, bold_font_name) for PDF export.

    ReportLab's built-in Helvetica does not reliably support symbols like
    INR/WON. We register DejaVu Sans from the local matplotlib install, which
    is typically available in this environment and has broad Unicode coverage.
    If anything goes wrong, we gracefully fall back to Helvetica.
    """
    try:
        from matplotlib import font_manager
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        regular_path = font_manager.findfont("DejaVu Sans", fallback_to_default=True)
        bold_path = font_manager.findfont("DejaVu Sans:style=Bold", fallback_to_default=True)

        regular_name = "SalaryScopeDejaVuSans"
        bold_name = "SalaryScopeDejaVuSansBold"

        registered = {font_name.lower() for font_name in pdfmetrics.getRegisteredFontNames()}
        if regular_name.lower() not in registered:
            pdfmetrics.registerFont(TTFont(regular_name, regular_path))
        if bold_name.lower() not in registered:
            pdfmetrics.registerFont(TTFont(bold_name, bold_path))

        return regular_name, bold_name
    except Exception:
        return "Helvetica", "Helvetica-Bold"


def _make_sections_pdf_bytes(
    report_title: str,
    summary_lines: tuple[str, ...],
    sections: list[tuple[str, pd.DataFrame]],
    chart_df: pd.DataFrame | None = None,
    chart_x: str | None = None,
    chart_y: str | None = None,
    chart_title: str | None = None,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    max_cols = max((len(df.columns) for _, df in sections), default=1)
    page_size = landscape(letter) if max_cols >= 7 else letter
    available_width = (10.0 if max_cols >= 7 else 7.0) * inch

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    body_style = styles["BodyText"].clone("CurrencyBodyCell")
    regular_font_name, bold_font_name = _get_pdf_font_names()
    body_style.fontName = regular_font_name
    body_style.fontSize = 8
    body_style.leading = 10
    body_style.wordWrap = "CJK"

    header_style = styles["BodyText"].clone("CurrencyHeaderCell")
    header_style.fontName = bold_font_name
    header_style.fontSize = 8
    header_style.leading = 10
    header_style.textColor = colors.white
    header_style.wordWrap = "CJK"

    title_style = styles["Title"].clone("CurrencyTitle")
    title_style.fontName = bold_font_name
    heading_style = styles["Heading2"].clone("CurrencyHeading2")
    heading_style.fontName = bold_font_name
    plain_style = styles["BodyText"].clone("CurrencyBody")
    plain_style.fontName = regular_font_name

    story = [Paragraph(report_title, title_style), Spacer(1, 10)]
    for line in summary_lines:
        if line:
            story.append(Paragraph(str(line), plain_style))
    if summary_lines:
        story.append(Spacer(1, 10))

    if chart_df is not None and chart_x and chart_y and not chart_df.empty:
        try:
            import matplotlib

            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            fig, ax = plt.subplots(figsize=(7.2, 3.8))
            ax.bar(chart_df[chart_x].astype(str), chart_df[chart_y].astype(float), color="#1A4F8A")
            ax.set_title(chart_title or "")
            ax.set_xlabel(chart_x)
            ax.set_ylabel(chart_y)
            ax.grid(axis="y", linestyle="--", alpha=0.35)
            ax.spines["top"].set_visible(False)
            ax.spines["right"].set_visible(False)
            plt.xticks(rotation=20)
            plt.tight_layout()

            img_buffer = BytesIO()
            fig.savefig(img_buffer, format="png", dpi=140, facecolor="white")
            plt.close(fig)
            img_buffer.seek(0)
            story.append(Image(img_buffer, width=6.5 * inch, height=3.4 * inch))
            story.append(Spacer(1, 10))
        except Exception:
            pass

    for idx, (section_title, df) in enumerate(sections):
        section_df = _clean_df(df)
        story.append(Paragraph(section_title, heading_style))
        story.append(Spacer(1, 6))

        string_df = section_df.astype(str)
        char_weights: list[int] = []
        for col in string_df.columns:
            values = [str(col)] + string_df[col].tolist()[: min(len(string_df), 12)]
            longest = max((len(v) for v in values), default=8)
            char_weights.append(min(max(longest, 8), 34))

        total_weight = max(sum(char_weights), 1)
        col_widths = [(weight / total_weight) * available_width for weight in char_weights]
        min_width = 0.8 * inch if len(section_df.columns) >= 7 else 1.0 * inch
        col_widths = [max(width, min_width) for width in col_widths]
        width_total = sum(col_widths)
        if width_total > available_width:
            scale = available_width / width_total
            col_widths = [width * scale for width in col_widths]

        header_row = [Paragraph(str(col), header_style) for col in section_df.columns]
        body_rows = [
            [Paragraph(str(value), body_style) for value in row]
            for row in string_df.values.tolist()
        ]

        table = Table([header_row] + body_rows, repeatRows=1, colWidths=col_widths)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A4F8A")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C7D2E0")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FB")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)
        if idx < len(sections) - 1:
            story.append(PageBreak())

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def _make_sections_docx_bytes(
    report_title: str,
    summary_lines: tuple[str, ...],
    sections: list[tuple[str, pd.DataFrame]],
) -> bytes | None:
    try:
        from docx import Document
    except ImportError:
        return None

    document = Document()
    document.add_heading(report_title, level=0)

    for line in summary_lines:
        if line:
            document.add_paragraph(str(line))

    for section_title, df in sections:
        document.add_heading(section_title, level=1)
        section_df = _clean_df(df)
        table = document.add_table(rows=1, cols=len(section_df.columns))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for idx, col in enumerate(section_df.columns):
            header_cells[idx].text = str(col)
        for _, row in section_df.iterrows():
            cells = table.add_row().cells
            for idx, value in enumerate(row.tolist()):
                cells[idx].text = str(value)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _prediction_sections(
    result_data: dict,
    target_currency: str,
    rate_data: dict,
    extra_sections: list[tuple[str, pd.DataFrame]] | None = None,
) -> tuple[tuple[str, ...], list[tuple[str, pd.DataFrame]]]:
    input_details = result_data.get("input_details") or result_data.get("input_details_a2") or {}
    prediction = float(result_data.get("prediction", result_data.get("prediction_a2", 0.0)))
    lower_bound = result_data.get("lower_bound")
    upper_bound = result_data.get("upper_bound")

    converted_annual = _convert_amount(prediction, target_currency, rate_data) or prediction
    converted_monthly = converted_annual / 12
    converted_weekly = converted_annual / 52
    converted_hourly = converted_annual / (52 * 40)

    summary_rows = {
        "Annual Salary": _format_currency(converted_annual, target_currency),
        "Monthly (Approx)": _format_currency(converted_monthly, target_currency),
        "Weekly (Approx)": _format_currency(converted_weekly, target_currency),
        "Hourly (Approx, 40hr/week)": _format_currency(converted_hourly, target_currency),
    }
    if lower_bound is not None and upper_bound is not None:
        lower_converted = _convert_amount(lower_bound, target_currency, rate_data)
        upper_converted = _convert_amount(upper_bound, target_currency, rate_data)
        summary_rows["Likely Salary Range"] = (
            f"{_format_currency(lower_converted, target_currency)} to "
            f"{_format_currency(upper_converted, target_currency)}"
        )
    if result_data.get("salary_band_label"):
        summary_rows["Salary Level"] = result_data["salary_band_label"]
    if result_data.get("career_stage_label"):
        summary_rows["Career Stage"] = result_data["career_stage_label"]
    if result_data.get("resume_score_data"):
        summary_rows["Resume Score"] = f"{result_data['resume_score_data'].get('total_score', 0)}/100"
    if result_data.get("resume_score_data_a2"):
        summary_rows["Resume Score"] = f"{result_data['resume_score_data_a2'].get('total_score_a2', 0)}/100"

    sections = [
        ("Input Details", pd.DataFrame(list(input_details.items()), columns=["Field", "Value"])),
        ("Prediction Summary", pd.DataFrame(list(summary_rows.items()), columns=["Metric", "Value"])),
    ]

    score_data = result_data.get("resume_score_data")
    if score_data:
        sections.append(
            (
                "Resume Score Breakdown",
                pd.DataFrame(
                    [
                        ("Total Score", f"{score_data.get('total_score', 0)}/100"),
                        ("Experience", score_data.get("experience_score", "")),
                        ("Education", score_data.get("education_score", "")),
                        ("Skills", score_data.get("skills_score", "")),
                        ("Profile Strength", score_data.get("level", "")),
                    ],
                    columns=["Metric", "Value"],
                ),
            )
        )
    score_data_a2 = result_data.get("resume_score_data_a2")
    if score_data_a2:
        sections.append(
            (
                "Resume Score Breakdown",
                pd.DataFrame(
                    [
                        ("Total Score", f"{score_data_a2.get('total_score_a2', 0)}/100"),
                        ("Experience", score_data_a2.get("experience_score_a2", "")),
                        ("Skills", score_data_a2.get("skills_score_a2", "")),
                        ("Role Relevance", score_data_a2.get("title_score_a2", "")),
                        ("Profile Strength", score_data_a2.get("level_a2", "")),
                    ],
                    columns=["Metric", "Value"],
                ),
            )
        )

    assoc_text = result_data.get("assoc_text_a1_improved")
    if assoc_text:
        sections.append(("Pattern Insight", pd.DataFrame([{"Insight": assoc_text}])))

    if extra_sections:
        sections.extend(extra_sections)

    return _summary_header_lines(target_currency, rate_data), sections


def _scenario_sections(
    results_df: pd.DataFrame,
    target_currency: str,
    rate_data: dict,
) -> tuple[tuple[str, ...], list[tuple[str, pd.DataFrame]], pd.DataFrame]:
    salary_col = "Predicted Salary (USD)"
    converted_col = f"Predicted Salary ({target_currency})"
    chart_df = results_df.copy()
    chart_df[converted_col] = chart_df[salary_col].astype(float).apply(
        lambda value: _convert_amount(value, target_currency, rate_data)
    )

    export_df = results_df.copy()
    export_df[converted_col] = chart_df[converted_col].round(2)
    export_df = export_df.drop(columns=[salary_col])

    if "Lower Bound" in export_df.columns and "Upper Bound" in export_df.columns:
        export_df[f"Lower Bound ({target_currency})"] = export_df["Lower Bound"].astype(float).apply(
            lambda value: _convert_amount(value, target_currency, rate_data)
        ).round(2)
        export_df[f"Upper Bound ({target_currency})"] = export_df["Upper Bound"].astype(float).apply(
            lambda value: _convert_amount(value, target_currency, rate_data)
        ).round(2)
        export_df = export_df.drop(columns=["Lower Bound", "Upper Bound"])

    avg_salary = chart_df[converted_col].mean()
    min_salary = chart_df[converted_col].min()
    max_salary = chart_df[converted_col].max()

    summary_lines = _summary_header_lines(target_currency, rate_data) + (
        f"Total scenarios: {len(export_df)}",
        f"Average predicted salary: {_format_currency(avg_salary, target_currency)}",
        f"Minimum predicted salary: {_format_currency(min_salary, target_currency)}",
        f"Maximum predicted salary: {_format_currency(max_salary, target_currency)}",
    )
    sections: list[tuple[str, pd.DataFrame]] = []

    summary_df = pd.DataFrame(
        [
            ("Total Scenarios", len(export_df)),
            ("Average Predicted Salary", _format_currency(avg_salary, target_currency)),
            ("Minimum Predicted Salary", _format_currency(min_salary, target_currency)),
            ("Maximum Predicted Salary", _format_currency(max_salary, target_currency)),
        ],
        columns=["Metric", "Value"],
    )
    sections.append(("Summary Statistics", summary_df))
    sections.append(("Scenario Results", export_df))

    if "Salary Level" in results_df.columns:
        level_group = (
            chart_df.groupby("Salary Level")[converted_col]
            .mean()
            .reset_index()
            .rename(columns={converted_col: f"Average Salary ({target_currency})"})
        )
        level_group[f"Average Salary ({target_currency})"] = level_group[f"Average Salary ({target_currency})"].round(2)
        sections.append(("Average Salary by Salary Level", level_group))

    if "Career Stage" in results_df.columns:
        stage_group = (
            chart_df.groupby("Career Stage")[converted_col]
            .mean()
            .reset_index()
            .rename(columns={converted_col: f"Average Salary ({target_currency})"})
        )
        stage_group[f"Average Salary ({target_currency})"] = stage_group[f"Average Salary ({target_currency})"].round(2)
        sections.append(("Average Salary by Career Stage", stage_group))

    if "Experience Level" in results_df.columns:
        exp_group = (
            chart_df.groupby("Experience Level")[converted_col]
            .mean()
            .reset_index()
            .rename(columns={converted_col: f"Average Salary ({target_currency})"})
        )
        exp_group[f"Average Salary ({target_currency})"] = exp_group[f"Average Salary ({target_currency})"].round(2)
        sections.append(("Average Salary by Experience Level", exp_group))

    if "Company Size" in results_df.columns:
        size_group = (
            chart_df.groupby("Company Size")[converted_col]
            .mean()
            .reset_index()
            .rename(columns={converted_col: f"Average Salary ({target_currency})"})
        )
        size_group[f"Average Salary ({target_currency})"] = size_group[f"Average Salary ({target_currency})"].round(2)
        sections.append(("Average Salary by Company Size", size_group))

    if "Work Mode" in results_df.columns:
        work_mode_group = (
            chart_df.groupby("Work Mode")[converted_col]
            .mean()
            .reset_index()
            .sort_values(converted_col, ascending=False)
            .rename(columns={converted_col: f"Average Salary ({target_currency})"})
        )
        work_mode_group[f"Average Salary ({target_currency})"] = work_mode_group[f"Average Salary ({target_currency})"].round(2)
        sections.append(("Average Salary by Work Mode", work_mode_group))

    best_row = chart_df.loc[chart_df[converted_col].idxmax()]
    worst_row = chart_df.loc[chart_df[converted_col].idxmin()]
    spread = float(best_row[converted_col] - worst_row[converted_col])
    insight_text = (
        f"Highest scenario: {best_row['Scenario']} at {_format_currency(float(best_row[converted_col]), target_currency)}. "
        f"Lowest scenario: {worst_row['Scenario']} at {_format_currency(float(worst_row[converted_col]), target_currency)}. "
        f"Spread: {_format_currency(spread, target_currency)}."
    )
    sections.append(("Insights", pd.DataFrame([{"Insight": insight_text}])))
    return summary_lines, sections, chart_df


@st.cache_data(show_spinner=False)
def build_manual_prediction_export_pdf(
    result_data: dict,
    target_currency: str,
    rate_data: dict,
    report_label: str,
    extra_sections: list[tuple[str, pd.DataFrame]] | None = None,
) -> bytes:
    summary_lines, sections = _prediction_sections(result_data, target_currency, rate_data, extra_sections=extra_sections)
    return _make_sections_pdf_bytes(
        f"SalaryScope {report_label} Report",
        summary_lines,
        sections,
    )


@st.cache_data(show_spinner=False)
def build_manual_prediction_export_docx(
    result_data: dict,
    target_currency: str,
    rate_data: dict,
    report_label: str,
    extra_sections: list[tuple[str, pd.DataFrame]] | None = None,
) -> bytes | None:
    summary_lines, sections = _prediction_sections(result_data, target_currency, rate_data, extra_sections=extra_sections)
    return _make_sections_docx_bytes(
        f"SalaryScope {report_label} Report",
        summary_lines,
        sections,
    )


@st.cache_data(show_spinner=False)
def build_resume_export_pdf(
    result_data: dict,
    target_currency: str,
    rate_data: dict,
    report_label: str,
    extra_sections: list[tuple[str, pd.DataFrame]] | None = None,
) -> bytes:
    summary_lines, sections = _prediction_sections(result_data, target_currency, rate_data, extra_sections=extra_sections)
    return _make_sections_pdf_bytes(
        f"SalaryScope {report_label} Report",
        summary_lines,
        sections,
    )


@st.cache_data(show_spinner=False)
def build_resume_export_docx(
    result_data: dict,
    target_currency: str,
    rate_data: dict,
    report_label: str,
    extra_sections: list[tuple[str, pd.DataFrame]] | None = None,
) -> bytes | None:
    summary_lines, sections = _prediction_sections(result_data, target_currency, rate_data, extra_sections=extra_sections)
    return _make_sections_docx_bytes(
        f"SalaryScope {report_label} Report",
        summary_lines,
        sections,
    )


@st.cache_data(show_spinner=False)
def build_scenario_export_pdf(results_df: pd.DataFrame, target_currency: str, rate_data: dict, report_label: str) -> bytes:
    summary_lines, sections, chart_df = _scenario_sections(results_df, target_currency, rate_data)
    converted_col = f"Predicted Salary ({target_currency})"
    return _make_sections_pdf_bytes(
        f"SalaryScope {report_label} Report",
        summary_lines,
        sections,
        chart_df=chart_df,
        chart_x="Scenario",
        chart_y=converted_col,
        chart_title=f"Predicted Salary Comparison ({target_currency})",
    )


@st.cache_data(show_spinner=False)
def build_scenario_export_docx(results_df: pd.DataFrame, target_currency: str, rate_data: dict, report_label: str) -> bytes | None:
    summary_lines, sections, _ = _scenario_sections(results_df, target_currency, rate_data)
    return _make_sections_docx_bytes(
        f"SalaryScope {report_label} Report",
        summary_lines,
        sections,
    )
