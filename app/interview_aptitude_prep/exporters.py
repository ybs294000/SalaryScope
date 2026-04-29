from __future__ import annotations

from io import BytesIO
from typing import Any

import pandas as pd
import streamlit as st

from .timer import format_duration


def _summary_lines(entry: dict[str, Any], result: dict[str, Any]) -> tuple[str, ...]:
    timer_snapshot = result.get("timer_snapshot", {})
    timed_label = "Timed" if timer_snapshot.get("timer_enabled") else "Untimed"
    elapsed_label = format_duration(timer_snapshot.get("elapsed_seconds"))
    return (
        f"Set: {result.get('title', entry.get('title', 'Practice Set'))}",
        f"Category: {entry.get('category', 'General')}",
        f"Role Focus: {entry.get('role_family', 'General')}",
        f"Difficulty: {entry.get('difficulty', result.get('difficulty', 'General'))}",
        f"Questions: {result.get('total_questions', 0)}",
        f"Attempt Mode: {timed_label}",
        f"Time Used: {elapsed_label}",
        f"Score: {result.get('total_obtained', 0):.2f} / {result.get('total_marks', 0):.2f}",
        f"Percent: {result.get('percentage', 0):.1f}%",
        f"Outcome: {'Pass' if result.get('percentage', 0) >= result.get('pass_percent', 0) else 'Review Needed'}",
    )


def _build_section_df(result: dict[str, Any]) -> pd.DataFrame:
    section_rows = []
    for section in result.get("section_summary", []):
        section_rows.append(
            {
                "Section": section.get("section_title", ""),
                "Score": round(float(section.get("obtained", 0.0)), 2),
                "Total Marks": round(float(section.get("marks", 0.0)), 2),
                "Correct": int(section.get("correct", 0)),
                "Skipped": int(section.get("skipped", 0)),
                "Questions": int(section.get("questions", 0)),
            }
        )
    return pd.DataFrame(section_rows)


def _build_question_df(result: dict[str, Any], review_settings: dict[str, Any]) -> pd.DataFrame:
    include_explanations = bool(review_settings.get("show_explanations", True))
    include_answers = bool(review_settings.get("show_correct_answers", True))

    question_rows = []
    for question in result.get("question_results", []):
        row = {
            "Section": question.get("section_title", ""),
            "Topic": question.get("topic", ""),
            "Prompt": question.get("prompt", ""),
            "Your Answer": question.get("user_answer_display", ""),
            "Result": "Correct" if question.get("is_correct") else ("Skipped" if question.get("is_skipped") else "Incorrect"),
            "Marks": round(float(question.get("marks", 0.0)), 2),
            "Obtained": round(float(question.get("obtained", 0.0)), 2),
        }
        if include_answers:
            row["Correct Answer"] = question.get("correct_answer_display", "")
        if include_explanations:
            row["Explanation"] = question.get("explanation", "")
        question_rows.append(row)
    return pd.DataFrame(question_rows)


@st.cache_data(show_spinner=False)
def build_interview_prep_csv(result: dict[str, Any], review_settings: dict[str, Any]) -> bytes:
    question_df = _build_question_df(result, review_settings)
    return question_df.to_csv(index=False).encode("utf-8")


@st.cache_data(show_spinner=False)
def build_interview_prep_pdf(
    entry: dict[str, Any],
    result: dict[str, Any],
    review_settings: dict[str, Any],
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    summary_lines = _summary_lines(entry, result)
    section_df = _build_section_df(result).fillna("")
    question_df = _build_question_df(result, review_settings).fillna("")
    page_size = landscape(letter) if len(question_df.columns) >= 7 else letter
    available_width = 10.0 * 72 if page_size == landscape(letter) else 7.0 * 72

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )

    styles = getSampleStyleSheet()
    story = [Paragraph("SalaryScope Interview Prep Report", styles["Title"]), Spacer(1, 10)]
    for line in summary_lines:
        story.append(Paragraph(str(line), styles["BodyText"]))
    story.append(Spacer(1, 10))

    if not section_df.empty:
        story.append(Paragraph("Section Summary", styles["Heading2"]))
        story.append(Spacer(1, 6))
        story.append(_make_pdf_table(section_df, available_width=available_width))
        story.append(Spacer(1, 10))

    if not question_df.empty:
        story.append(Paragraph("Answer Review", styles["Heading2"]))
        story.append(Spacer(1, 6))
        story.append(_make_pdf_table(question_df, max_col_chars=52, available_width=available_width))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


@st.cache_data(show_spinner=False)
def build_interview_prep_docx(
    entry: dict[str, Any],
    result: dict[str, Any],
    review_settings: dict[str, Any],
) -> bytes | None:
    try:
        from docx import Document
    except ImportError:
        return None

    summary_lines = _summary_lines(entry, result)
    section_df = _build_section_df(result).fillna("")
    question_df = _build_question_df(result, review_settings).fillna("")

    document = Document()
    document.add_heading("SalaryScope Interview Prep Report", level=0)
    for line in summary_lines:
        document.add_paragraph(str(line))

    if not section_df.empty:
        document.add_heading("Section Summary", level=1)
        _append_docx_table(document, section_df)

    if not question_df.empty:
        document.add_heading("Answer Review", level=1)
        _append_docx_table(document, question_df)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _make_pdf_table(df: pd.DataFrame, max_col_chars: int = 28, available_width: float = 7.0 * 72) -> Table:
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, Table, TableStyle

    styles = getSampleStyleSheet()
    header_style = styles["BodyText"].clone("IAPDFHeader")
    header_style.fontName = "Helvetica-Bold"
    header_style.fontSize = 8
    header_style.leading = 10
    header_style.textColor = colors.white
    header_style.wordWrap = "CJK"

    body_style = styles["BodyText"].clone("IAPDFBody")
    body_style.fontName = "Helvetica"
    body_style.fontSize = 7.5
    body_style.leading = 9.5
    body_style.wordWrap = "CJK"

    clipped = df.astype(str).copy()
    for col in clipped.columns:
        clipped[col] = clipped[col].map(lambda value: value if len(value) <= max_col_chars else value[: max_col_chars - 3] + "...")

    header_row = [Paragraph(str(col), header_style) for col in clipped.columns]
    body_rows = [
        [Paragraph(str(value), body_style) for value in row]
        for row in clipped.values.tolist()
    ]

    data = [header_row] + body_rows
    col_count = max(len(clipped.columns), 1)
    col_width = available_width / col_count
    table = Table(data, repeatRows=1, colWidths=[col_width] * col_count)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A4F8A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7D2E0")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FB")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _append_docx_table(document: Any, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        header_cells[idx].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row.tolist()):
            cells[idx].text = str(value)
