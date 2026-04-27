"""
hr_tools/export_utils.py
------------------------
Shared export helpers for HR & Employer Tools.

Why a separate file instead of extending app/utils/pdf_utils.py?
- HR exports are mostly compact table/report outputs, not the large multi-page
  prediction reports handled by pdf_utils.py.
- Keeping them here keeps the HR tool package modular and easy to remove.
- It also keeps future HR-specific formatting changes isolated from the main
  prediction-report pipeline.
"""

from __future__ import annotations

from io import BytesIO
from typing import Iterable

import pandas as pd
import streamlit as st


def _clean_df_for_export(df: pd.DataFrame, max_rows: int | None = None) -> pd.DataFrame:
    cleaned = df.copy()
    if max_rows is not None:
        cleaned = cleaned.head(max_rows)
    cleaned = cleaned.fillna("")

    for col in cleaned.columns:
        if pd.api.types.is_float_dtype(cleaned[col]):
            cleaned[col] = cleaned[col].round(2)
    return cleaned


def _safe_sheet_name(name: str) -> str:
    invalid = ["\\", "/", "*", "[", "]", ":", "?"]
    safe = str(name)
    for ch in invalid:
        safe = safe.replace(ch, "-")
    return safe[:31] or "Sheet1"


@st.cache_data(show_spinner=False)
def build_xlsx_bytes(sections: list[tuple[str, pd.DataFrame]]) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for title, df in sections:
            _clean_df_for_export(df).to_excel(
                writer,
                sheet_name=_safe_sheet_name(title),
                index=False,
            )
    buffer.seek(0)
    return buffer.getvalue()


@st.cache_data(show_spinner=False)
def build_pdf_bytes(
    report_title: str,
    sections: list[tuple[str, pd.DataFrame]],
    summary_lines: tuple[str, ...] = (),
    max_rows_per_section: int = 40,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    max_cols = max((len(df.columns) for _, df in sections), default=1)
    page_size = landscape(letter) if max_cols >= 8 else letter
    available_width = (10.0 if max_cols >= 8 else 7.0) * inch

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    body_style = styles["BodyText"].clone("HRBodyCell")
    body_style.fontName = "Helvetica"
    body_style.fontSize = 7.5 if max_cols >= 8 else 8.2
    body_style.leading = 9.2 if max_cols >= 8 else 10
    body_style.wordWrap = "CJK"

    header_style = styles["BodyText"].clone("HRHeaderCell")
    header_style.fontName = "Helvetica-Bold"
    header_style.fontSize = 7.5 if max_cols >= 8 else 8.2
    header_style.leading = 9.2 if max_cols >= 8 else 10
    header_style.textColor = colors.white
    header_style.wordWrap = "CJK"

    story = []

    story.append(Paragraph(report_title, styles["Title"]))
    story.append(Spacer(1, 10))
    for line in summary_lines:
        if line:
            story.append(Paragraph(str(line), styles["BodyText"]))
    if summary_lines:
        story.append(Spacer(1, 10))

    for idx, (section_title, df) in enumerate(sections):
        section_df = _clean_df_for_export(df, max_rows=max_rows_per_section)
        story.append(Paragraph(section_title, styles["Heading2"]))
        story.append(Spacer(1, 6))

        string_df = section_df.astype(str)
        char_weights: list[int] = []
        for col in string_df.columns:
            values = [str(col)] + string_df[col].tolist()[: min(len(string_df), 12)]
            longest = max((len(v) for v in values), default=8)
            char_weights.append(min(max(longest, 8), 34))

        total_weight = max(sum(char_weights), 1)
        col_widths = [(weight / total_weight) * available_width for weight in char_weights]
        min_width = 0.7 * inch if len(section_df.columns) >= 8 else 0.85 * inch
        max_width = 2.0 * inch if len(section_df.columns) >= 8 else 2.4 * inch
        col_widths = [min(max(width, min_width), max_width) for width in col_widths]

        width_total = sum(col_widths)
        if width_total > available_width:
            scale = available_width / width_total
            col_widths = [width * scale for width in col_widths]

        header_row = [Paragraph(str(col), header_style) for col in section_df.columns]
        body_rows = [
            [Paragraph(str(value), body_style) for value in row]
            for row in string_df.values.tolist()
        ]

        data = [header_row] + body_rows
        table = Table(data, repeatRows=1, colWidths=col_widths)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A4F8A")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C7D2E0")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FB")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)

        if len(df) > len(section_df):
            story.append(Spacer(1, 6))
            story.append(
                Paragraph(
                    f"Showing the first {len(section_df)} rows of {len(df)} in this section.",
                    styles["Italic"],
                )
            )

        if idx < len(sections) - 1:
            story.append(PageBreak())

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


@st.cache_data(show_spinner=False)
def build_docx_bytes(
    report_title: str,
    sections: list[tuple[str, pd.DataFrame]],
    summary_lines: tuple[str, ...] = (),
    max_rows_per_section: int = 50,
) -> bytes | None:
    try:
        from docx import Document
    except ImportError:
        return None

    document = Document()
    document.add_heading(report_title, level=0)

    for line in summary_lines:
        if line:
            document.add_paragraph(str(line))

    for section_title, df in sections:
        document.add_heading(section_title, level=1)
        section_df = _clean_df_for_export(df, max_rows=max_rows_per_section)

        table = document.add_table(rows=1, cols=len(section_df.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, col in enumerate(section_df.columns):
            hdr[idx].text = str(col)

        for _, row in section_df.iterrows():
            cells = table.add_row().cells
            for idx, value in enumerate(row.tolist()):
                cells[idx].text = str(value)

        if len(df) > len(section_df):
            document.add_paragraph(
                f"Showing the first {len(section_df)} rows of {len(df)} in this section."
            )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def render_export_buttons(
    *,
    title: str,
    file_stem: str,
    csv_df: pd.DataFrame,
    xlsx_sections: list[tuple[str, pd.DataFrame]] | None = None,
    pdf_sections: list[tuple[str, pd.DataFrame]] | None = None,
    docx_sections: list[tuple[str, pd.DataFrame]] | None = None,
    summary_lines: Iterable[str] = (),
    key_prefix: str,
    csv_label: str = "Download CSV",
    xlsx_label: str = "Download XLSX",
    pdf_label: str = "Download PDF",
    docx_label: str = "Download DOCX",
) -> None:
    summary_tuple = tuple(str(x) for x in summary_lines if str(x).strip())
    xlsx_sections = xlsx_sections or [("Data", csv_df)]
    pdf_sections = pdf_sections or [("Data", csv_df)]
    docx_sections = docx_sections or pdf_sections

    csv_bytes = _clean_df_for_export(csv_df).to_csv(index=False).encode("utf-8")
    xlsx_bytes = build_xlsx_bytes(xlsx_sections)
    pdf_bytes = build_pdf_bytes(title, pdf_sections, summary_tuple)
    docx_bytes = build_docx_bytes(title, docx_sections, summary_tuple)

    st.markdown("#### Export Results")
    st.caption("Download the current result in the format that best fits your workflow.")

    cols = st.columns(4 if docx_bytes is not None else 3)
    with cols[0]:
        st.download_button(
            csv_label,
            data=csv_bytes,
            file_name=f"{file_stem}.csv",
            mime="text/csv",
            key=f"{key_prefix}_csv",
            width="stretch",
        )
    with cols[1]:
        st.download_button(
            xlsx_label,
            data=xlsx_bytes,
            file_name=f"{file_stem}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"{key_prefix}_xlsx",
            width="stretch",
        )
    with cols[2]:
        st.download_button(
            pdf_label,
            data=pdf_bytes,
            file_name=f"{file_stem}.pdf",
            mime="application/pdf",
            key=f"{key_prefix}_pdf",
            width="stretch",
        )

    if docx_bytes is not None:
        with cols[3]:
            st.download_button(
                docx_label,
                data=docx_bytes,
                file_name=f"{file_stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"{key_prefix}_docx",
                width="stretch",
            )
